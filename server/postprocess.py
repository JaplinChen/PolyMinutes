"""Post-meeting pass over the recording.

The live pipeline trades accuracy for latency: a small model, online clustering that cannot see
what comes later, and one-shot refinement. None of those constraints apply once the meeting ends,
so this re-runs the whole thing from the wav with the largest model available and clusters over
every segment at once, then rewrites the stored transcript.
"""

from __future__ import annotations

import contextlib
import json
import logging
import os
from contextlib import AbstractContextManager
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

import numpy as np
import soundfile as sf

from . import asr, asr_gpu, config, correct, diarize, jobs, translate
from .store import Store

log = logging.getLogger("polyminutes.postprocess")

# Utterances a speaker must have before their own language statistics outweigh the meeting's.
MIN_LANGUAGE_EVIDENCE = 4
# A speaker whose language disagrees with the meeting's needs this share of the meeting's lines
# before that disagreement is believed. Four lines is enough to settle zh against zh-yue; it is not
# enough to declare an English speaker in a Chinese meeting, because the clusters that do that are
# not people. Measured on a 2h19m Mandarin factory meeting: three speakers held nothing but
# English — 5, 9 and 11 lines of "Thank you for watching" and "Each police officer" — and each
# cleared the flat bar of four, certified itself English, and was never re-decoded. A participant
# who really does speak English says more than 1.8% of a meeting.
MIN_MINORITY_SHARE = 0.05
# Utterances handed to the recogniser at once. Long enough to amortise the encoder, short enough
# that an interrupted run has reported most of what it did.
BATCH_UTTERANCES = 64
# The shortest piece a speaker change is allowed to carve out of an utterance. Below this it is
# someone agreeing mid-sentence, and cutting there costs a transcript line to gain nothing.
MIN_PIECE_SECONDS = 0.5
# The shortest utterance the post-meeting pass will hand to the recogniser. The live path keeps the
# VAD's own 0.25 because a one-word confirmation belongs on the TV; here the opposite is true.
# Whisper never declines — handed a fragment with no sentence in it, it writes the YouTube subtitle
# sign-offs it read most, and a real meeting came back with "剪輯 李宗盛" four times. Measured on a
# 2h19m factory meeting, VAD then turn-splitting, against seventeen known sign-off lines:
#
#     min speech   utterances   speech kept   sign-offs surviving   long real lines kept
#           0.25          958      105.8 min             15 of 17              6 of 6
#           0.50          894      103.5 min             11 of 17              6 of 6
#           0.80          824       99.7 min              8 of 17              6 of 6
#           1.00          743       94.7 min              7 of 17              5 of 6
#
# 0.80 was chosen off that table as the last value costing no long line. Re-measured on
# 2026-08-21 over two whole meetings, decoded end to end at each setting, it was too high — the
# table counted sign-offs, and what 0.80 was also buying was silence:
#
#     min speech   transcribed speech (s3 / s6)   gained vs 0.80   lost vs 0.80
#           0.25            30.4 / 120.1 min       4.6 / 10.0 min   0.2 / 1.5 min
#           0.50            29.1 / 117.4 min       3.3 /  7.0 min   0.1 / 1.2 min
#           0.80            26.0 / 111.6 min             —               —
#
# 0.50 gains 3.3 and 7.0 minutes of real speech and gives back almost nothing; the 11 added runs
# over three seconds on each meeting are whole reports (今天早上由品管來報告一下我們七月份第五週
# 內銷客戶的質量表現期…, 總經理還有問一下現場的幹部…). 0.25 gains another 1.3 and 2.7 minutes but
# doubles the sub-1.5s fragments, which is where the sign-offs the table was counting live.
#
# What that table could not see, because it counted VAD utterances rather than decoding them:
# the threshold also moves what lands in a decode batch, and with it the one language faster-
# whisper reports for the whole batch. Session 6 came back with 64 runs of English translation at
# 0.80 and none at 0.50 — same audio, same weights. That is a batch-language problem, not a VAD
# one; the setting must not be read as a fix for it.
POST_MEETING_MIN_SPEECH = 0.50
# Silence one speaker may leave inside a single decoded clip. Above it they have stopped; below
# it they are breathing, and cutting there hands Whisper a fragment with no sentence in it.
#
# Swept on the 2026-08-05 meeting, after diarize.regroup_speakers made the labels worth honouring.
# The chairman's eleven-minute address is the section that matters — it is where the fragments and
# the hallucinations both were:
#
#     gap    lines   median   under 3s   chairman median
#     1.5      231     3.9 s        101            2.4 s
#     2.0      208     5.0 s         74            3.7 s
#     2.5      189     6.0 s         63            3.9 s
#     3.0      177     7.0 s         53            4.9 s
#
# 2.0 is where honouring the speaker labels catches up with ignoring them entirely — the same
# recording merged on the gap alone, labels discarded, gives 208 lines and a 4.0 s chairman
# median. Past that the merge is joining runs the labels do not endorse, which is the attribution
# error this whole path exists to avoid.
MERGE_MAX_GAP = 2.0
# The longest clip the merge will build. Whisper's window is 30s and it discards the rest without
# saying so, so this leaves headroom; it also stays under segment.py's MAX_MERGED_SECONDS, which
# is what decides whether a transcript line is still scrubbable.
MERGE_MAX_SECONDS = 20.0

# Utterances averaged into a speaker's stored voiceprint: enough for a stable centroid, and a
# fixed cost per speaker rather than one embedding per utterance in the meeting.
VOICEPRINT_SAMPLES = 5
# The utterance length a voiceprint sample wants. Same lesson as the naming clips (#145/#146):
# the *longest* utterance is adversely selected — a line is long exactly when the segmenter missed
# a turn inside it, so it averages somebody else's voice into the centroid. Mid-monologue pieces
# near this length are the clean ones.
VOICEPRINT_IDEAL_SECONDS = 8.0


@dataclass
class Utterance:
    start: float
    samples: np.ndarray
    speaker: str = ""
    lang: str = ""
    text: str = ""
    # Which decode this utterance's language came out of. Utterances sharing one are not
    # independent readings of what language was spoken — see `dominant_languages`.
    decode: int = -1


def best_model() -> Path:
    """Largest Whisper tier present on disk. Accuracy matters here, speed does not."""
    available = config.available_whisper_models()
    if not available:
        raise FileNotFoundError(f"no Whisper model found under {config.MODELS_DIR}")
    order = list(config.WHISPER_DIRS)
    return config.WHISPER_DIRS[max(available, key=order.index)]


def segment(wav: Path) -> list[Utterance]:
    audio, rate = sf.read(str(wav), dtype="float32")
    if rate != config.SAMPLE_RATE:
        raise ValueError(f"{wav} is {rate} Hz, expected {config.SAMPLE_RATE}")
    if audio.ndim > 1:
        audio = audio.mean(axis=1)

    vad = asr.Vad(min_speech=POST_MEETING_MIN_SPEECH)
    out: list[Utterance] = []
    for i in range(0, len(audio), config.BLOCK_SIZE):
        out += [Utterance(s.start, s.samples) for s in vad.push(audio[i : i + config.BLOCK_SIZE])]
    out += [Utterance(s.start, s.samples) for s in vad.flush()]
    return out


def split_on_turns(utterances: list[Utterance], turns: list[diarize.Turn]) -> list[Utterance]:
    """Cut every utterance where the speaker changes inside it, and label each piece.

    A VAD utterance is speech between silences, which is not the same thing as one person talking:
    on a 2h19m meeting 72 of 859 held more than one voice, one of them four people over twelve
    seconds. Those became a single averaged embedding and a single transcript line reading as
    nonsense, because it was four people's sentences run together.

    Pieces below MIN_PIECE_SECONDS are folded into the neighbour they touch rather than kept: a
    half-second of someone agreeing in the middle of a sentence is not a turn worth a line.
    """
    if not turns:
        return utterances

    out: list[Utterance] = []
    for u in utterances:
        seconds = len(u.samples) / config.SAMPLE_RATE
        cuts = _cut_points(u.start, u.start + seconds, turns)
        if len(cuts) < 2:
            u.speaker = diarize.speaker_code(_dominant(u.start, u.start + seconds, turns))
            out.append(u)
            continue
        for begin, end, who in cuts:
            head = int(round((begin - u.start) * config.SAMPLE_RATE))
            tail = int(round((end - u.start) * config.SAMPLE_RATE))
            out.append(Utterance(begin, u.samples[head:tail], speaker=diarize.speaker_code(who)))

    _inherit_missing(out)
    return out


def _inherit_missing(utterances: list[Utterance]) -> None:
    """Give the segmenter's blind spots to whoever was talking around them.

    A half-second the VAD called speech and the segmentation model called nobody would otherwise
    reach the transcript with a blank where the speaker goes. Five of 1040 pieces on a real
    meeting, all under a second. Same rule the clustering path applies to a clip too short to
    embed: it belongs to the person already speaking.
    """
    previous = ""
    for u in utterances:
        if u.speaker:
            previous = u.speaker
        else:
            u.speaker = previous
    # Anything before the first labelled piece has nothing behind it to inherit from.
    following = ""
    for u in reversed(utterances):
        if u.speaker:
            following = u.speaker
        else:
            u.speaker = following


def merge_runs(utterances: list[Utterance], wav: Path) -> list[Utterance]:
    """Rejoin what one speaker said without stopping, so Whisper decodes sentences and not breaths.

    The VAD ends an utterance on silence, which is not where a sentence ends. On the 2026-08-05
    factory meeting that left 157 of 258 lines under three seconds — and the short lines are where
    the transcript falls apart. Same recording, same model, same hour:

        section                 lines   median length   median text
        the general manager        87           7.3 s      27 chars
        the chairman              125           1.8 s      10 chars

    The chairman's section is the one that came back reading 「料理 互動」and「MICROPHONE SELFIE」.
    Nothing is wrong with the audio there — he pauses more, so the VAD cuts more, so Whisper sees
    1.8 seconds with no sentence in it and writes what such audio sounded like in its training set.
    Decoding the same seconds as one 20-second clip gives it the context it needs to decline.

    Runs are built after the speaker split, so a run is one person by construction: anyone else
    taking the floor lands between them in the list and ends the run. The clip is sliced from the
    recording rather than concatenated from the pieces, which keeps the pauses the speaker actually
    left — the timestamps stay honest, and Whisper is given silence rather than a splice.

    segment.py still runs later and still merges at the text level; this does not replace it. It
    fixes the half segment.py cannot reach, because by then the hallucination is already written.
    """
    if len(utterances) < 2:
        return utterances

    audio, rate = sf.read(str(wav), dtype="float32")
    if rate != config.SAMPLE_RATE:
        raise ValueError(f"{wav} is {rate} Hz, expected {config.SAMPLE_RATE}")
    if audio.ndim > 1:
        audio = audio.mean(axis=1)

    def ends(u: Utterance) -> float:
        return u.start + len(u.samples) / config.SAMPLE_RATE

    def flush(run: list[Utterance]) -> Utterance:
        if len(run) == 1:
            return run[0]
        head = int(round(run[0].start * config.SAMPLE_RATE))
        tail = int(round(ends(run[-1]) * config.SAMPLE_RATE))
        return Utterance(run[0].start, audio[head:tail], speaker=run[0].speaker)

    out: list[Utterance] = []
    run: list[Utterance] = [utterances[0]]
    for u in utterances[1:]:
        gap = u.start - ends(run[-1])
        span = ends(u) - run[0].start
        if u.speaker == run[0].speaker and 0 <= gap <= MERGE_MAX_GAP and span <= MERGE_MAX_SECONDS:
            run.append(u)
            continue
        out.append(flush(run))
        run = [u]
    out.append(flush(run))
    log.info("merged %d utterances into %d runs", len(utterances), len(out))
    return out


def _join(pieces: list[tuple[float, float, int]]) -> list[tuple[float, float, int]]:
    """Neighbouring pieces on the same speaker are one piece."""
    joined: list[tuple[float, float, int]] = []
    for begin, end, who in pieces:
        if joined and joined[-1][2] == who:
            joined[-1] = (joined[-1][0], end, who)
        else:
            joined.append((begin, end, who))
    return joined


def _fill_blind(pieces: list[tuple[float, float, int]]) -> list[tuple[float, float, int]]:
    """Stretches the segmenter heard nobody in belong to whoever was talking around them.

    Cutting on every turn edge exposes the pauses between one person's own turns, which would
    otherwise reach the transcript as a blank-speaker line carved out of the middle of their
    sentence. An utterance the segmenter heard nobody in at all keeps its -1 for `_inherit_missing`
    to resolve against the utterances either side of it.
    """
    known = [who for _, _, who in pieces if who >= 0]
    if not known:
        return pieces
    filled, previous = [], known[0]
    for begin, end, who in pieces:
        previous = who if who >= 0 else previous
        filled.append((begin, end, previous))
    return filled


def _cut_points(start: float, end: float, turns: list[diarize.Turn]) -> list[tuple[float, float, int]]:
    """The utterance split into (start, end, speaker), or one span when nobody else speaks in it.

    Cut on every turn edge, then let the shortest turn covering each stretch own it. Walking the
    turns in order and cutting at each one's end — the previous rule — could only see turns queued
    one after another: a turn nested inside a longer one ends behind the point already reached, so
    it was skipped and the second voice stayed in the line. That is what overlapping speech looks
    like coming out of the segmenter, and it is not rare: on a real 2.7h meeting 33 of 918 turns
    were nested, leaving 12 transcript lines (95 seconds) holding a voice the segmenter had marked.
    """
    edges = sorted({start, end} | {e for t in turns for e in (t.start, t.end) if start < e < end})
    pieces = _fill_blind([(a, b, _dominant(a, b, turns)) for a, b in zip(edges, edges[1:])])
    pieces = _join(pieces)
    # Shortest first, because absorbing one piece can leave its neighbour the shortest in turn.
    while len(pieces) > 1:
        i = min(range(len(pieces)), key=lambda i: pieces[i][1] - pieces[i][0])
        if pieces[i][1] - pieces[i][0] >= MIN_PIECE_SECONDS:
            break
        # Into the longer neighbour: a half-second of someone agreeing mid-sentence is not a turn
        # worth a line, and the sentence it interrupts is the one it belongs to.
        before = pieces[i - 1] if i else None
        after = pieces[i + 1] if i + 1 < len(pieces) else None
        into = before if after is None or (before and before[1] - before[0] >= after[1] - after[0]) \
            else after
        pieces[i] = (pieces[i][0], pieces[i][1], into[2])
        pieces = _join(pieces)
    return pieces or [(start, end, _dominant(start, end, turns))]


def _dominant(start: float, end: float, turns: list[diarize.Turn]) -> int:
    """Whoever holds most of this span. -1 when the segmenter heard nobody in it.

    Ties go to whoever started talking most recently. Between two turn edges every covering turn
    holds the whole stretch, so a tie is the normal case there and something has to break it: the
    newest voice owns the overlap. Both shapes of overlapping speech then cut — the turn nested
    inside a longer one wins the seconds it was marked for, and the turn that begins before the
    outgoing speaker has finished takes the handover with it. Preferring the *shorter* turn instead
    gets the nested case right and the handover wrong: the outgoing speaker's turn is the shorter
    one there, so the line stayed whole. Measured on a 2.7h meeting: 3 of the 12 uncut lines.
    """
    best, best_key = -1, (0.0, 0.0)
    for turn in turns:
        key = (min(turn.end, end) - max(turn.start, start), turn.start)
        if key[0] > 0 and key > best_key:
            best, best_key = turn.speaker, key
    return best


def _remember_voices(store: Store, session_id: int, utterances: list[Utterance],
                     diarizer: diarize.Diarizer) -> dict[str, str]:
    """Store one voiceprint per speaker, and recognise the ones the room already knows.

    Only the live pipeline stored voiceprints, which meant an imported recording never taught
    anything: naming S3 on the transcript page looked up a voiceprint that had never been written,
    found nothing, and quietly skipped promoting it. The whole point of the naming screen is that
    the next meeting recognises the voice.

    Returns the codes a known voice was recognised on, so a reprocess can put the learned names
    back rather than dropping the room to anonymous Sn again.
    """
    groups: dict[str, list[Utterance]] = {}
    # A piece flanked by the same speaker on both sides is mid-monologue; one at a speaker
    # boundary is where a missed turn leaves the other person's voice. Prefer the former.
    boundary: set[int] = set()
    for i, u in enumerate(utterances):
        prev = utterances[i - 1].speaker if i > 0 else ""
        nxt = utterances[i + 1].speaker if i + 1 < len(utterances) else ""
        if prev != u.speaker or nxt != u.speaker:
            boundary.add(id(u))
        if u.speaker and len(u.samples) / config.SAMPLE_RATE >= config.MIN_EMBED_SECONDS:
            groups.setdefault(u.speaker, []).append(u)

    recognised: dict[str, str] = {}
    for code, said in groups.items():
        # A clean handful, not everything they said: a centroid is an average, and averaging the
        # cleanest samples costs a fixed few embeddings per speaker instead of one per utterance.
        best = sorted(said, key=lambda u: (
            id(u) in boundary,
            abs(len(u.samples) / config.SAMPLE_RATE - VOICEPRINT_IDEAL_SECONDS),
        ))[:VOICEPRINT_SAMPLES]
        centroid = np.mean([diarizer.embed(u.samples) for u in best], axis=0).astype(np.float32)
        store.save_voiceprint(session_id, code, centroid.tobytes())
        name, score = diarizer.recognise(centroid)
        if name:
            recognised[code] = name
            # A confident match is labelled data nobody had to type: fold this meeting's print
            # back into the name so a drifting voice keeps refreshing its own variants.
            if score >= config.AUTO_LEARN_THRESHOLD:
                store.remember_speaker(name, centroid.tobytes())
    return recognised


def assign_speakers(utterances: list[Utterance], diarizer: diarize.Diarizer) -> None:
    """Cluster over the whole meeting at once.

    The fallback for a machine without the segmentation model. Clusters whole VAD utterances, so
    an utterance holding two people is one embedding averaging both.

    This is what the online pass could not do: a speaker whose first few seconds were atypical
    gets merged with their later segments instead of living on as a phantom second participant.
    """
    # Same rule the live path applies in Diarizer.assign: a clip too short to embed reliably
    # inherits the previous speaker. Clustering it instead mints a phantom participant per blip —
    # on a real recording the ten idle minutes before the meeting produced fourteen of them.
    long = [u for u in utterances if len(u.samples) / config.SAMPLE_RATE >= config.MIN_EMBED_SECONDS]
    if not long:
        for u in utterances:
            u.speaker = "S1"
        return

    labels = diarize.cluster_offline([diarizer.embed(u.samples) for u in long])
    for utterance, label in zip(long, labels):
        utterance.speaker = f"S{label + 1}"

    previous = f"S{labels[0] + 1}"
    for u in utterances:
        if u.speaker:
            previous = u.speaker
        else:
            u.speaker = previous


def dominant_languages(utterances: list[Utterance]) -> dict[str, str]:
    """Majority language per speaker, computed after clustering rather than as the meeting ran.

    A speaker needs to have said enough for a majority to mean anything. Raising the clustering
    threshold to separate real participants also produces a long tail of speakers holding two or
    three utterances, and letting those establish their own language put 433 Chinese lines under
    an English label across seven interviews — 產品 產品 產品 decoded as English because one
    stray detection was all the evidence there was.

    Below the minimum a speaker inherits the meeting's language, which is a far better guess than
    a coin flip on two samples.
    """
    # One vote per decode, not per utterance. The batched recogniser lays sixty-four clips end to
    # end and faster-whisper reports one language for the whole run, which it then hands back
    # against every clip — so counting them individually invents sixty-four agreeing witnesses out
    # of one. Measured on the 2026-08-10 meeting: one batch came back English and put 64 runs of
    # English translation into a Mandarin factory meeting (「This is the electric fan that is
    # often reminded of」 for 這個是常常提醒的那個電風扇), 8.4% of the transcript — comfortably
    # over MIN_MINORITY_SHARE, which is sized for a real minority speaker and cannot be lowered to
    # catch this without discarding one. Counted once, that batch is a single dissenting vote
    # against a dozen, and every clip in it is re-decoded in the speaker's own language.
    #
    # The unbatched path leaves `decode` at -1 and is given a distinct id per utterance here, so
    # a recogniser that really did read each clip separately keeps every vote it earned.
    seen: set[tuple[str, str, int]] = set()
    counts: dict[str, dict[str, int]] = {}
    overall: dict[str, int] = {}
    for i, u in enumerate(utterances):
        # Text-less utterances are dropped noise; their detected language is Whisper guessing at
        # static and must not vote.
        if not (u.lang and u.text):
            continue
        ballot = (u.speaker, u.lang, u.decode if u.decode >= 0 else ~i)
        if ballot in seen:
            continue
        seen.add(ballot)
        counts.setdefault(u.speaker, {})[u.lang] = counts.setdefault(u.speaker, {}).get(u.lang, 0) + 1
        overall[u.lang] = overall.get(u.lang, 0) + 1

    if not overall:
        return {}
    meeting = max(overall, key=overall.get)
    floor = sum(overall.values()) * MIN_MINORITY_SHARE
    return {code: _majority(langs, meeting, floor) for code, langs in counts.items() if langs}


def _majority(langs: dict[str, int], meeting: str, floor: float) -> str:
    best = max(langs, key=langs.get)
    if sum(langs.values()) < MIN_LANGUAGE_EVIDENCE:
        return meeting
    if best != meeting and langs[best] < floor:
        return meeting
    return best


def transcribe_all(utterances: list[Utterance], transcriber: asr.Transcriber,
                   progress: Callable[[Utterance, int, int], None] | None = None,
                   forced: dict[str, str] | None = None) -> None:
    """Two passes: detect each speaker's language, then re-transcribe anyone who was decoded
    under a language that disagrees with their majority.

    `forced` maps a speaker code to the language the room has set for that recognised voice. It wins
    over the detected majority: the majority is a guess from auto-detect, which flips a Chinese
    speaker to Vietnamese often enough that a voice the room has already identified should not be
    left to it. An empty or absent entry falls back to the majority, unchanged.

    `progress` is called after each first-pass decode. A ninety-minute recording spends most of an
    hour in that first loop, and a caller with somewhere to put partial results should not have to
    wait for the whole thing to survive an interruption.
    """
    # The first pass is the expensive one and every utterance in it wants the same thing —
    # auto-detect — so it goes through the recogniser in batches when the recogniser has a batch
    # mode. Boundaries are preserved either way; only the number of round trips changes.
    batch = getattr(transcriber, "transcribe_many", None)
    if batch:
        done = 0
        for start in range(0, len(utterances), BATCH_UTTERANCES):
            group = utterances[start : start + BATCH_UTTERANCES]
            for u, (text, lang) in zip(group, batch([g.samples for g in group], "")):
                u.text, u.lang, u.decode = text, lang, start
                done += 1
                if progress:
                    progress(u, done, len(utterances))
    else:
        for i, u in enumerate(utterances, 1):
            u.text, u.lang = transcriber.transcribe(u.samples, "")
            if progress:
                progress(u, i, len(utterances))

    dominant = dominant_languages(utterances)
    forced = forced or {}
    for u in utterances:
        want = forced.get(u.speaker) or dominant.get(u.speaker, "")
        # An utterance the first pass refused is retried too, not only one decoded under the wrong
        # language. Auto-detect collapses on perfectly ordinary speech often enough that dropping
        # those outright cost 992 real lines across seven interviews — 掃描機這件事情有 and
        # 就會直接進到系統變成需求 among them. The speaker's own language usually recovers them.
        if want and (not u.text or want != u.lang):
            text, used = transcriber.transcribe(u.samples, want)
            # Empty again, and the glossary prompt is the last thing left to remove: it is a prior
            # the decoder can be talked out of a sentence by, and the recogniser that has one
            # offers a run without it. 5 of 33 utterances lost on the 2026-08-10 meeting came back
            # only here.
            if not text and (plain := getattr(transcriber, "transcribe_unbiased", None)):
                text, used = plain(u.samples, want)
            # Still empty means the speaker's own language decoded this as noise as well, which is
            # what static sounds like to Whisper. Drop it rather than keep a phantom line.
            u.text, u.lang = (text, used) if text else ("", u.lang)


def translated_rows(entries, store: Store, cfg: config.Config,
                    translator: translate.Translator | None,
                    stop: Callable[[], bool]) -> list[dict]:
    """Corrected, translated line rows from (start, end_time, speaker, lang, text) entries.

    One loop for both ways a transcript comes to exist — decoded from audio or read off a subtitle
    track — so the corrector, the rolling context, the glossary and the failed-translation status
    behave identically whichever produced the text.
    """
    terms = store.glossary()
    corrector = correct.Corrector(terms, store.corrections())
    context: list[translate.Line] = []
    rows: list[dict] = []
    for start, end_time, speaker, lang, text in entries:
        if stop():
            raise jobs.Cancelled()
        if not text:
            continue
        text = corrector.fix(text)
        line = translate.Line(text=text, lang=lang, speaker=speaker)
        targets = [c for c in cfg.languages if c != lang]
        translations: dict[str, str] = {}
        status = "ok"
        if translator and targets:
            try:
                translations = translator.translate(
                    line, targets, context=context[-config.CONTEXT_LINES:], terms=terms
                ).translations
            except Exception:
                # Recorded rather than swallowed: a line with no translation and a line whose
                # translation failed look identical in the transcript, and only one of them is
                # worth re-running.
                log.exception("translation failed at %.2fs", start)
                status = "translate_failed"
        rows.append({
            "start": start,
            "end_time": end_time,
            "speaker": speaker,
            "lang": lang,
            "source": text,
            "translations": translations,
            "status": status,
        })
        context.append(line)
    return rows


def subtitle_session(store: Store, session_id: int, cues: list[tuple[float, float, str]],
                     lang: str, cfg: config.Config,
                     translator: translate.Translator | None = None,
                     should_stop: Callable[[], bool] | None = None) -> list[dict]:
    """Store a transcript taken from a subtitle track: no decode, no card, one unknown speaker.

    A subtitle track carries no voices, so every line lands on S1 — reprocess runs the full
    pipeline from the audio if the speakers matter. Everything else is the shared row loop, so the
    lines read like any other transcript's.
    """
    stop = should_stop or (lambda: False)
    rows = translated_rows(((s, e, "S1", lang, t) for s, e, t in cues), store, cfg,
                           translator, stop)
    if not rows:
        raise ValueError("the subtitle track had no usable cues")
    store.replace_lines(session_id, rows)
    return rows


def rewrite_session(store: Store, session_id: int, wav: Path, cfg: config.Config,
                    translator: translate.Translator | None = None,
                    should_stop: Callable[[], bool] | None = None,
                    gpu: Callable[[], AbstractContextManager] = contextlib.nullcontext,
                    ) -> list[Utterance]:
    """Re-derive the transcript and replace the stored lines for this session.

    `should_stop` lets a meeting starting in the room take the GPU back. It is polled between
    decode batches and between translations, and acting on it costs nothing: the stored transcript
    is untouched until `replace_lines` at the very end, so an abandoned pass leaves the session
    exactly as it found it.

    `gpu` guards the one stage that uses the card. Only decoding does: the VAD, the speaker
    segmentation and the per-line translation round trips all run on the CPU or the network, and
    they are most of the wall clock — on a 2h19m recording the segmentation alone is eight
    minutes. Holding the card across them means someone pressing record waits for all of it, and
    `claim_gpu` gives up after thirty seconds.

    The caller passes the guard rather than this function taking one, because the two callers need
    different behaviour and neither can be inferred from here: an import queues for the card
    (`borrow_gpu`), while the pass scheduled after a meeting must yield to the next one. Taking a
    gate here as well as in the caller would deadlock — `jobs.py`'s gate is a semaphore of one and
    the scheduled path waits on it without a timeout.
    """
    stop = should_stop or (lambda: False)
    # GPU first. The CPU fallback keeps float32 weights and every core: this runs after the
    # meeting, so accuracy is the only concern — but it also makes the machine unusable while it
    # runs, which is the other reason the GPU path exists.
    transcriber = (asr_gpu.maybe(cfg.languages, asr_gpu.hotwords_from(store.glossary()))
                   or asr.Transcriber(model_dir=best_model(), quantized=False,
                                      num_threads=os.cpu_count() or 4, languages=cfg.languages))
    # Known voices loaded so a re-derive can put the room's learned names back on the freshly
    # clustered codes — a reprocess renumbers speakers from scratch, and without this every name the
    # user typed would land on a different person or none at all.
    diarizer = diarize.Diarizer(cfg=cfg, known=diarize.load_known(store))

    utterances = segment(wav)
    log.info("%d utterances from %s", len(utterances), wav.name)
    if not utterances:
        return []

    # Speaker turns where the model for them is installed, clustered whole utterances where it is
    # not. The first also decides the transcript's line boundaries, because a line that runs across
    # a speaker change belongs to neither of them.
    speaker_turns = diarize.turns(wav)
    if speaker_turns:
        # Before anything is labelled: the turn pass splits one long speaker across many labels,
        # and every stage after this — the line's speaker, the voiceprints, whether merge_runs can
        # join two of his sentences — reads those labels as identity.
        speaker_turns = diarize.regroup_speakers(wav, speaker_turns, diarizer.embed)
        before = len(utterances)
        utterances = split_on_turns(utterances, speaker_turns)
        log.info("%d turns split %d utterances into %d", len(speaker_turns), before, len(utterances))
    else:
        assign_speakers(utterances, diarizer)
    recognised = _remember_voices(store, session_id, utterances, diarizer)

    # After the voiceprints, deliberately: they are picked by utterance length (VOICEPRINT_IDEAL_
    # SECONDS), and merging first would hand that selection a pool of clips that are all long for a
    # different reason. From here on the merged runs are the transcript's lines.
    utterances = merge_runs(utterances, wav)

    def watch(_u: Utterance, _done: int, _total: int) -> None:
        if stop():
            raise jobs.Cancelled()

    # A recognised voice with a language set is transcribed in it, not in whatever auto-detect
    # guessed — the room's Vietnamese speaker stays Vietnamese, its Chinese speakers stay Chinese.
    languages = store.speaker_languages()
    forced = {code: lang for code, name in (recognised or {}).items() if (lang := languages.get(name))}

    with gpu():
        transcribe_all(utterances, transcriber, progress=watch, forced=forced)

    # The stored transcript is not touched until every line is translated. Replacing it line by
    # line as they came in meant a failure halfway through left the session holding half a
    # transcript, and a failure right after the delete left it holding none.
    entries = ((u.start, u.start + len(u.samples) / config.SAMPLE_RATE, u.speaker, u.lang, u.text)
               for u in utterances)
    rows = translated_rows(entries, store, cfg, translator, stop)

    # replace_lines deletes the old transcript before inserting the new. An empty result would make
    # that a bare delete — and a re-transcription that decoded nothing (every utterance came back
    # empty without raising) must not wipe a transcript that already exists, the same way the
    # no-utterances case above returns without touching it. A session that never had lines is still
    # allowed to be created empty, as a silent import is.
    if not rows and store.lines(session_id):
        raise ValueError("re-transcription produced no lines; keeping the existing transcript")
    store.replace_lines(session_id, rows)
    # Only now that the new transcript is committed: the old names were keyed to the old codes, which
    # this pass renumbered, so they are cleared and replaced by whatever the learned voices matched.
    # Done here, not beside _remember_voices, so a cancelled pass leaves the names as it found them.
    store.clear_speaker_names(session_id)
    for code, name in recognised.items():
        store.set_speaker_name(session_id, code, name)
    return utterances


def _summary_markdown(store: Store, session_id: int, names: dict[str, str]) -> list[str]:
    """The summary block of the export, every generated language in full.

    The export's reader is whoever was not in the room, so unlike the page there is no interface
    language to pick by — all of them go in. A summary generated before the transcript's latest
    edit is still included, but says so: an exported file that silently carried outdated decisions
    would be trusted precisely because it is a file.
    """
    row = store.summary(session_id)
    if not row or row["status"] == "failed":
        return []
    try:
        per_language = json.loads(row["json"])
    except ValueError:
        return []

    session = store.session(session_id)
    stale = bool(session) and int(session["lines_rev"]) != int(row["lines_rev"])

    out = ["## 會議摘要", ""]
    if stale:
        out += ["> ⚠ 摘要生成後逐字稿曾被修改，內容可能與下方逐字稿不一致。", ""]
    for lang, s in per_language.items():
        out += [f"### {s.get('title') or lang}（{lang}）", "", s.get("summary", ""), ""]
        if s.get("decisions"):
            out += ["**決議**", ""] + [f"- {d}" for d in s["decisions"]] + [""]
        if s.get("actions"):
            out += ["**行動項目**", ""]
            out += [f"- {names.get(a.get('speaker', ''), a.get('speaker')) or '（未指定）'}："
                    f"{a.get('text', '')}" for a in s["actions"]]
            out += [""]
    return out


def to_markdown(store: Store, session_id: int) -> str:
    """Speaker-attributed transcript with every language stacked under each turn."""
    lines = store.lines(session_id)
    names = store.speaker_names(session_id)
    if not lines:
        return "# 會議紀錄\n\n（無內容）\n"

    out = ["# 會議紀錄", ""]
    out += _summary_markdown(store, session_id, names)
    # Numeric, not lexicographic: codes are S1..S35, and a plain string sort put S10 between S1 and
    # S2. The tuple key keeps any non-S<n> code (there should be none) after the numbered ones
    # without an int-vs-str comparison error.
    speakers = sorted({l["speaker"] for l in lines},
                      key=lambda s: (0, int(s[1:])) if s[1:].isdigit() else (1, s))
    out += ["## 發言者", ""]
    out += [f"- **{names.get(code, code)}**" + ("" if code in names else "（未命名）") for code in speakers]
    out += ["", "## 逐字稿", ""]

    for line in lines:
        stamp = f"{int(line['start']) // 60}:{int(line['start']) % 60:02d}"
        who = names.get(line["speaker"], line["speaker"])
        out.append(f"**[{stamp}] {who}**")
        out.append(f"> {line['source']}")
        for lang, text in line["translations"].items():
            out.append(f"> _{lang}_ {text}")
        out.append("")

    return "\n".join(out) + "\n"


def to_docx(store: Store, session_id: int) -> bytes:
    """The same export as `to_markdown`, as a Word document.

    Enterprise meetings hand over a .docx, not a code fence. It carries the same three blocks in the
    same order — summary, speakers, transcript — and the same honesty: a summary older than the
    latest transcript edit is included but flagged, because a file is trusted precisely for being a
    file. python-docx is imported here, not at module load, so the recognition path never pays for a
    dependency only the export uses.
    """
    from docx import Document  # noqa: PLC0415 — export-only, kept off the hot import path

    lines = store.lines(session_id)
    names = store.speaker_names(session_id)
    doc = Document()
    doc.add_heading("會議紀錄", level=0)
    if not lines:
        doc.add_paragraph("（無內容）")
        return _docx_bytes(doc)

    row = store.summary(session_id)
    per_language = None
    if row and row["status"] != "failed":
        try:
            per_language = json.loads(row["json"])
        except ValueError:
            per_language = None
    if per_language:
        session = store.session(session_id)
        stale = bool(session) and int(session["lines_rev"]) != int(row["lines_rev"])
        doc.add_heading("會議摘要", level=1)
        if stale:
            doc.add_paragraph("⚠ 摘要生成後逐字稿曾被修改，內容可能與下方逐字稿不一致。")
        for lang, s in per_language.items():
            doc.add_heading(f"{s.get('title') or lang}（{lang}）", level=2)
            if s.get("summary"):
                doc.add_paragraph(s["summary"])
            if s.get("decisions"):
                doc.add_paragraph("決議").bold = True
                for d in s["decisions"]:
                    doc.add_paragraph(str(d), style="List Bullet")
            if s.get("actions"):
                doc.add_paragraph("行動項目").bold = True
                for a in s["actions"]:
                    who = names.get(a.get("speaker", ""), a.get("speaker")) or "（未指定）"
                    doc.add_paragraph(f"{who}：{a.get('text', '')}", style="List Bullet")

    speakers = sorted({l["speaker"] for l in lines},
                      key=lambda s: (0, int(s[1:])) if s[1:].isdigit() else (1, s))
    doc.add_heading("發言者", level=1)
    for code in speakers:
        doc.add_paragraph(names.get(code, code) + ("" if code in names else "（未命名）"),
                          style="List Bullet")

    doc.add_heading("逐字稿", level=1)
    for line in lines:
        stamp = f"{int(line['start']) // 60}:{int(line['start']) % 60:02d}"
        who = names.get(line["speaker"], line["speaker"])
        head = doc.add_paragraph()
        head.add_run(f"[{stamp}] {who}").bold = True
        doc.add_paragraph(line["source"])
        for lang, text in line["translations"].items():
            doc.add_paragraph(f"{lang}｜{text}")

    return _docx_bytes(doc)


def _docx_bytes(doc) -> bytes:
    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _vtt_escape(text: str) -> str:
    """Cue text is markup: an unescaped `<` starts a tag and swallows the rest of the line."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _vtt_time(seconds: float) -> str:
    """HH:MM:SS.mmm. Rounded in milliseconds, not in seconds-plus-fraction: rounding the fraction
    of 3.9996 on its own carries to 1000 and writes `.1000`, which no player will parse."""
    ms = max(int(round(seconds * 1000)), 0)
    return f"{ms // 3600000:02d}:{ms // 60000 % 60:02d}:{ms // 1000 % 60:02d}.{ms % 1000:03d}"


def to_vtt(store: Store, session_id: int, lang: str | None = None) -> str:
    """The transcript as subtitles, one track per language.

    A meeting already plays back beside its video (#170); WebVTT is what makes that transcript
    usable in any other player too. One language per file rather than every language stacked in one
    cue, because a player shows a cue verbatim and three stacked languages cover the picture — as
    separate tracks the viewer picks one. `lang=None` is the spoken source.
    """
    lines = store.lines(session_id)
    names = store.speaker_names(session_id)

    cues: list[tuple[float, float, str, str]] = []
    for line in lines:
        text = line["source"] if lang is None else line["translations"].get(lang, "")
        if not (text or "").strip():
            continue
        start = float(line["start"])
        # end_time predates most of the corpus and is still nullable; the same two-second fallback
        # the clip queries use keeps an old meeting from emitting zero-length cues players discard.
        end = float(line["end_time"]) if line["end_time"] is not None else start + 2.0
        cues.append((start, end, names.get(line["speaker"], line["speaker"]), text.strip()))

    out = ["WEBVTT", ""]
    for i, (start, end, who, text) in enumerate(cues):
        # A cue overlapping the next one makes players show both or jump; the recogniser's end
        # times routinely run past the next speaker's start when two people talk over each other.
        limit = cues[i + 1][0] if i + 1 < len(cues) else end
        end = max(min(end, limit), start + 0.1)
        out += [f"{_vtt_time(start)} --> {_vtt_time(end)}",
                f"<v {_vtt_escape(who)}>{_vtt_escape(text)}", ""]
    return "\n".join(out) + "\n"
